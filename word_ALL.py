import pandas as pd
from docx import Document
import json
path=r'Z:\Software\WebSphere MQ 6.0_Client\Licenses\Windows\project-x\lcdpfcd\L7DF227L72CN.docx'
data=Document(path)
pathout=r"Z:\Software\WebSphere MQ 6.0_Client\Licenses\Windows\mTool4\data\output.txt"
with open (pathout,'w') as f:
    f.write('')
patha=r'Z:\Software\WebSphere MQ 6.0_Client\Licenses\Windows\temp\pfcd_route.json'

with open(patha, 'r', encoding='utf-8') as f:
    mydic = json.load(f)

from_cell='L76F127K74X1'##change
new_cell= 'L76F127K7412'##change


real_from_cell=''##change新建不同偏貼 PFCD來源PFCD右邊或者是跨廠-->3

new_cell_unpzat='L76F127K74X2'##change第11碼X
new_cell_merge=''##change切割後 merge

tft_source='TJ6F12DK'##change
cf_source='FJ6F12XKR3FJ6F12XKR4'##change

tft_source_merge=''##change
cf_source_merge=''##change
if new_cell_unpzat!='':
    mytype=1 #主線產生未偏貼出貨,有SOURCE TFT CF
elif new_cell_merge!='':
    mytype=2 #新玻璃切割後併入
elif tft_source!='':
    mytype=0
elif real_from_cell!='':
    mytype=3#新建不同偏貼 PFCD,新包材導入
print("select PFCD,ROUTE_ID from L7WPT1D.CPRDCT")

myfrom=['* Start Stage： Before PI','* Start Stage： Before Polarizer','* Start Stage： Before PI','* Start Stage： Before Polarizer']
myto=['* Finish Stage：After LOT2','* Finish Stage：Before Polarizer','* Finish Stage：Before Polarizer','* Finish Stage：After LOT2']

mypzat=["偏光板	未偏貼出貨","未偏貼出貨","切割後 merge 回原 PFCD","偏光板	日東"]
mycfpzat=["","","","CF  4101B00667000-BV5M"]
mytftpzat=["","","","TFT 4101B00729000-DGH2"]
mycfpzatrwk=["","","",""]
mytftpzatrwk=["","","",""]
my_desc='DF39 Border 5.5+江蘇合成 新增日東 POL'

print(mypzat[mytype])
pathfrompfcd=r"Z:\Software\WebSphere MQ 6.0_Client\Licenses\Windows\mTool4\data\frompfcd.txt"
with open (pathfrompfcd,'w') as f:
    f.write(f'{from_cell}')

if mytype==1:
    pathnewpfcd=r"Z:\Software\WebSphere MQ 6.0_Client\Licenses\Windows\mTool4\data\newpfcd.txt"
    with open (pathnewpfcd,'w') as f:
        f.write(f'{new_cell_unpzat}')
else:
    pathnewpfcd=r"Z:\Software\WebSphere MQ 6.0_Client\Licenses\Windows\mTool4\data\newpfcd.txt"
    with open (pathnewpfcd,'w') as f:
        f.write(f'{new_cell}')


if from_cell[2:6]!=new_cell[2:6]:
    print("***********先將SPC CODE ITEM匯出,再匯入七個檔案*************")
    with open (pathout,'a') as f:
        print("***********先將SPC CODE ITEM匯出,再匯入七個檔案*************",file=f)
        ##f.write('"***********先將SPC CODE ITEM匯出,再匯入七個檔案*************\n')
    import os
    directory=r'Z:\Software\WebSphere MQ 6.0_Client\Licenses\Windows\mTool4\SQL\lcd_new_pfcd'
    directory1=r'Z:\Software\WebSphere MQ 6.0_Client\Licenses\Windows\mTool4\mTool4\data'

    def updateFile(file,old_str,new_str):
        file_data=""
        with open(file,'r') as f:
            for line in f:
                line=line.replace(old_str,new_str)
                file_data += line
        with open(file,'w') as f:
            f.write(file_data)

    path=[os.path.join(directory,d) for d in os.listdir(directory)]
    with open(path[0],'r') as f:
        mydata=f.read()
    myoldstr=mydata.split("%")[-2]

    for i in range(len(path)):

        updateFile(path[i],myoldstr,from_cell[2:6]) ##CHANGE
    path1=[os.path.join(directory1,d) for d in os.listdir(directory1)]
    for i in range(len(path1)):

        updateFile(path1[i],from_cell[2:6],new_cell[2:6])
##        updateFile(path1[i],'","64.2','","LCW') ##CHANGE



if new_cell_unpzat!='':
    print("==EXPORT PARAM兩個TABLE+SHIP QTIME==")
    print(f"pfcd in ('{from_cell}') order by ope_id")
    print(f"BRM建立PFCD : *{from_cell[2:6]}*{from_cell[7]}*{from_cell[9:]}*")
    print(f"'{new_cell_unpzat}',L{tft_source[1:]}")
    ##print(f"'{new_cell_merge}','{new_cell_merge[0]}J{new_cell_merge[2:]}','{new_cell_merge[0]}J{new_cell_merge[2:8]}X{new_cell_merge[9:]}','FJ{new_cell_merge[2:6]}X{new_cell_merge[7:]}',{tft_source_merge},{cf_source_merge}")

    print("==將SQL存成:pfcd_pid.txt==")
    with open (pathout,'a') as f:
        print(f"select distinct ope_id,recipe_id from L7WPT1D.CPARAM WHERE substr(pfcd,3,4)='{from_cell[2:6]}' and substr(pfcd,8,1)='{from_cell[7]}' \
and substr(pfcd,10,3)='{from_cell[9:]}' \
and ope_id in ('3000','4000','4100','4600','4800','3270') and recipe_id!='' ORDER BY OPE_ID",file=f)
elif new_cell_merge!='':
    print("==EXPORT PARAM兩個TABLE+SHIP QTIME==")
    print(f"pfcd in ('{from_cell[0]}J{from_cell[2:]}','{from_cell[0]}J{from_cell[2:8]}X{from_cell[9:]}','FJ{from_cell[2:6]}X{from_cell[7:]}') order by ope_id")
    print(f"BRM建立PFCD : *{from_cell[2:6]}*{from_cell[7]}*{from_cell[9:]}*")
    print(f"'{new_cell}','{new_cell[0]}J{new_cell[2:]}','{new_cell[0]}J{new_cell[2:8]}X{new_cell[9:]}','FJ{new_cell[2:6]}X{new_cell[7:]}',{tft_source},{cf_source},L{tft_source[1:]}")
    ##print(f"'{new_cell_merge}','{new_cell_merge[0]}J{new_cell_merge[2:]}','{new_cell_merge[0]}J{new_cell_merge[2:8]}X{new_cell_merge[9:]}','FJ{new_cell_merge[2:6]}X{new_cell_merge[7:]}',{tft_source_merge},{cf_source_merge}")

    print("==將SQL存成:pfcd_pid.txt==")
    with open (pathout,'a') as f:
        print(f"select distinct ope_id,recipe_id from L7WPT1D.CPARAM WHERE substr(pfcd,3,4)='{from_cell[2:6]}' and substr(pfcd,8,1)='{from_cell[7]}' \
and substr(pfcd,10,3)='{from_cell[9:]}' \
and ope_id in ('2100','2133','2200','2210','2700','3000','4000','4100','4600','4800','3270','4200') and eqpt_id not in ('GAPX0200','DUMP0100') and recipe_id!='' ORDER BY OPE_ID",file=f)
elif tft_source!='':
    print("==EXPORT PARAM兩個TABLE+SHIP QTIME==")
    print(f"pfcd in ('{from_cell}','{from_cell[0]}J{from_cell[2:]}','{from_cell[0]}J{from_cell[2:8]}X{from_cell[9:]}','FJ{from_cell[2:6]}X{from_cell[7:]}') order by ope_id")
    print(f"BRM建立PFCD : *{from_cell[2:6]}*{from_cell[7]}*{from_cell[9:]}*")
    print(f"'{new_cell}','{new_cell[0]}J{new_cell[2:]}','{new_cell[0]}J{new_cell[2:8]}X{new_cell[9:]}','FJ{new_cell[2:6]}X{new_cell[7:]}',{tft_source},{cf_source},L{tft_source[1:]}")
    ##print(f"'{new_cell_merge}','{new_cell_merge[0]}J{new_cell_merge[2:]}','{new_cell_merge[0]}J{new_cell_merge[2:8]}X{new_cell_merge[9:]}','FJ{new_cell_merge[2:6]}X{new_cell_merge[7:]}',{tft_source_merge},{cf_source_merge}")

    print("==將SQL存成:pfcd_pid.txt==")
    with open (pathout,'a') as f:
        print(f"select distinct ope_id,recipe_id from L7WPT1D.CPARAM WHERE substr(pfcd,3,4)='{from_cell[2:6]}' and substr(pfcd,8,1)='{from_cell[7]}' \
and substr(pfcd,10,3)='{from_cell[9:]}' \
and ope_id in ('2100','2133','2200','2210','2700','3000','4000','4100','4600','4800','3270','4200') and eqpt_id not in ('GAPX0200','DUMP0100') and recipe_id!='' ORDER BY OPE_ID",file=f)
else:
    print("==EXPORT PARAM兩個TABLE+SHIP QTIME==")
    print(f"pfcd in ('{from_cell}') order by ope_id")
    print(f"BRM建立PFCD : *{from_cell[2:6]}*{from_cell[7]}*{from_cell[9:]}*")
    print(f"'{new_cell}'")
    ##print(f"'{new_cell_merge}','{new_cell_merge[0]}J{new_cell_merge[2:]}','{new_cell_merge[0]}J{new_cell_merge[2:8]}X{new_cell_merge[9:]}','FJ{new_cell_merge[2:6]}X{new_cell_merge[7:]}',{tft_source_merge},{cf_source_merge}")

    print("==將SQL存成:pfcd_pid.txt==")
    with open (pathout,'a') as f:
        print(f"select distinct ope_id,recipe_id from L7WPT1D.CPARAM WHERE substr(pfcd,3,4)='{from_cell[2:6]}' and substr(pfcd,8,1)='{from_cell[7]}' \
and substr(pfcd,10,3)='{from_cell[9:]}' \
and ope_id in ('3000','4000','4050','4052','4053','4100','4600','4800','3270','4200') and recipe_id!='' and eqpt_id not in ('GAPX0200','DUMP0100') ORDER BY OPE_ID",file=f)
if new_cell_unpzat!='':
    data.tables[0].cell(0,2).paragraphs[0].text=f'*LCD PFCD：{new_cell_unpzat}'
elif new_cell_merge!='':
    data.tables[0].cell(0,2).paragraphs[0].text=f'*LCD PFCD：{new_cell_merge}'
else:
    data.tables[0].cell(0,2).paragraphs[0].text=f'*LCD PFCD：{new_cell}'
data.tables[0].cell(1,0).paragraphs[0].text=f'*Description：{my_desc}'
data.tables[0].cell(2,0).paragraphs[0].text='* Polarizer (主/替/替)   料號如下 :'
data.tables[0].cell(2,0).paragraphs[1].text=f'{mypzat[mytype]}'
data.tables[0].cell(2,0).paragraphs[2].text=f'{mycfpzat[mytype]}'
data.tables[0].cell(2,0).paragraphs[3].text=f'{mytftpzat[mytype]}'
data.tables[0].cell(2,0).paragraphs[4].text=f'{mycfpzatrwk[mytype]}'
data.tables[0].cell(2,0).paragraphs[5].text=f'{mytftpzatrwk[mytype]}'
data.tables[0].cell(3,0).paragraphs[0].text=f'{myfrom[mytype]}'
data.tables[0].cell(3,11).paragraphs[0].text=f'{myto[mytype]}'
data.tables[0].cell(13,0).paragraphs[0].text=f'PFCD ：  '
data.tables[0].cell(13,15).paragraphs[0].text=f'*Part No.：     '
data.tables[0].cell(13,20).paragraphs[3].text=' '
if tft_source!='':
    data.tables[0].cell(21,0).paragraphs[0].text=f'PFCD ： {new_cell[0]}J{new_cell[2:]} '#{new_cell_merge[0]}J{new_cell_merge[2:]}'
    a_route=new_cell[0]+'J'+new_cell[2:]
    data.tables[0].cell(21,0).paragraphs[1].text=f'Route：{mydic[a_route]}'
    data.tables[0].cell(21,7).paragraphs[0].text=f'PFCD ： {new_cell[0]}J{new_cell[2:8]}X{new_cell[9:]} '#{new_cell_merge[0]}J{new_cell_merge[2:8]}X{new_cell_merge[9:]}'
    b_route=new_cell[0]+'J'+new_cell[2:8]+'X'+new_cell[9:]
    data.tables[0].cell(21,7).paragraphs[1].text=f'Route：{mydic[b_route]}'
    data.tables[0].cell(22,7).paragraphs[0].text=f'PFCD ： FJ{new_cell[2:6]}X{new_cell[7:]} '#FJ{new_cell_merge[2:6]}X{new_cell_merge[7:]}'
    c_route='FJ'+new_cell[2:6]+'X'+new_cell[7:]
    data.tables[0].cell(22,7).paragraphs[1].text=f'Route：{mydic[c_route]}'
else:
    data.tables[0].cell(21,0).paragraphs[0].text=f'PFCD ：  '
    a_route=new_cell[0]+'J'+new_cell[2:]
    data.tables[0].cell(21,0).paragraphs[1].text=f'Route：'
    data.tables[0].cell(21,7).paragraphs[0].text=f'PFCD ：  '
    b_route=new_cell[0]+'J'+new_cell[2:8]+'X'+new_cell[9:]
    data.tables[0].cell(21,7).paragraphs[1].text=f'Route：'
    data.tables[0].cell(22,7).paragraphs[0].text=f'PFCD ： '
    c_route='FJ'+new_cell[2:6]+'X'+new_cell[7:]
    data.tables[0].cell(22,7).paragraphs[1].text=f'Route：'
    data.tables[0].cell(13,0).paragraphs[0].text=f'PFCD ： {new_cell} '
    data.tables[0].cell(13,15).paragraphs[0].text=f'*Part No.：    {real_from_cell} '
    if mycfpzat[mytype]=='':
        data.tables[0].cell(13,20).paragraphs[3].text='3000'
    elif mytftpzatrwk[mytype]=='':
        data.tables[0].cell(13,20).paragraphs[3].text='4100'
    else:
        data.tables[0].cell(13,20).paragraphs[3].text='4000'
if new_cell_unpzat!='':
    d_route=new_cell_unpzat
elif new_cell_merge!='':
    d_route=new_cell_merge
else:
    d_route=new_cell

data.tables[0].cell(9,6).paragraphs[0].text=f'* Route：{mydic[d_route]}'
data.tables[0].cell(21,15).paragraphs[0].text=f'{tft_source}'
data.tables[0].cell(22,15).paragraphs[0].text=f'{cf_source}'
data.tables[0].cell(21,15).paragraphs[1].text=f'{tft_source_merge}'
data.tables[0].cell(22,15).paragraphs[1].text=f'{cf_source_merge}'
if new_cell_unpzat!='':
    path1=rf'Z:\Software\WebSphere MQ 6.0_Client\Licenses\Windows\pfcd\{new_cell_unpzat}.docx'
elif new_cell_merge!='':
    path1=rf'Z:\Software\WebSphere MQ 6.0_Client\Licenses\Windows\pfcd\{new_cell_merge}.docx'
else:
    path1=rf'Z:\Software\WebSphere MQ 6.0_Client\Licenses\Windows\pfcd\{new_cell}.docx'

data.save(path1)
'''
patha=r'Z:\Software\WebSphere MQ 6.0_Client\Licenses\Windows\temp\pfcd_pid.txt'

dfpid=pd.read_fwf(patha)
##print(dfpid)
dfpid=dfpid.drop([0])
dfpid['OPE_ID']=dfpid['OPE_ID           RECIPE_ID'].apply(lambda x: x.split(' ')[0])
dfpid['RECIPE_ID']=dfpid['OPE_ID           RECIPE_ID'].apply(lambda x: x.split(' ')[-1])
dfpid=dfpid.drop(columns=['OPE_ID           RECIPE_ID'])
print(dfpid)
pathb=r'Z:\Software\WebSphere MQ 6.0_Client\Licenses\Windows\temp\pfcd_pid.csv'
dfpid.to_csv(pathb, index=False)
with open (pathout,'a') as f:
    f.write(f'{dfpid}')
'''